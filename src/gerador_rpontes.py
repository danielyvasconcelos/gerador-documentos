"""
Gerador de Documentos - R Pontes Construtora
Automatiza a criação de propostas comerciais a partir de dados do Excel
"""

import openpyxl
from docx import Document
import os
import time
from datetime import datetime
from validador import ValidadorDados
from logger import Logger

class GeradorRPontes:
    """Classe principal que gera documentos Word a partir de dados do Excel"""
    
    def __init__(self, caminho_planilha, caminho_template=None, pasta_saida=None):
        self.caminho_planilha = caminho_planilha
        self.caminho_template = caminho_template or "../templates/modelo_proposta.docx"
        self.pasta_saida = pasta_saida or "../data/documentos_gerados"
        self.logger = Logger()
        
        # Cria pasta de saída se não existir
        if not os.path.exists(self.pasta_saida):
            os.makedirs(self.pasta_saida)
        
    def ler_dados_completos(self):
        """Lê e combina dados das 4 abas da planilha RPONTES"""
        self.logger.info(f"Carregando planilha: {os.path.basename(self.caminho_planilha)}")
        wb = openpyxl.load_workbook(self.caminho_planilha)
        
        # Carrega cada aba da planilha
        aba_clientes = wb['Dados dos Clientes']
        aba_imoveis = wb['Dados do ImóvelServiço Contrata']
        aba_financeiro = wb['Dados Financeiros e Valores']
        aba_status = wb['StatusErro']
        
        dados_completos = []
        max_row = aba_clientes.max_row
        
        # Percorre cada linha de dados (pula o cabeçalho)
        for i in range(2, max_row + 1):
            cliente = {}
            
            # Dados pessoais
            cliente['ID_PROPOSTA'] = aba_clientes.cell(i, 1).value
            cliente['NOME_CLIENTE'] = aba_clientes.cell(i, 2).value
            cliente['DOCUMENTO_CLIENTE'] = aba_clientes.cell(i, 3).value
            cliente['EMAIL_CLIENTE'] = aba_clientes.cell(i, 4).value
            cliente['DATA_PROPOSTA'] = aba_clientes.cell(i, 5).value
            
            # Dados do imóvel
            cliente['EMPREENDIMENTO'] = aba_imoveis.cell(i, 1).value
            cliente['TIPO_UNIDADE'] = aba_imoveis.cell(i, 2).value
            cliente['METRAGEM'] = aba_imoveis.cell(i, 3).value
            cliente['DESCRICAO_SERVICOS'] = aba_imoveis.cell(i, 4).value
            
            # Dados financeiros - formata valores em reais
            valor_total = aba_financeiro.cell(i, 1).value
            valor_entrada = aba_financeiro.cell(i, 2).value
            cliente['VALOR_TOTAL'] = f"R$ {valor_total:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            cliente['VALOR_ENTRADA'] = f"R$ {valor_entrada:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
            cliente['NUMERO_PARCELAS'] = int(aba_financeiro.cell(i, 3).value)
            cliente['DATA_PRIMEIRA_PARCELA'] = aba_financeiro.cell(i, 4).value
            
            # Status de controle
            cliente['STATUS_GERACAO'] = aba_status.cell(i, 1).value
            
            dados_completos.append(cliente)
            
        wb.close()
        self.logger.info(f"Dados carregados: {len(dados_completos)} clientes encontrados")
        return dados_completos
    
    def gerar_documento(self, dados_cliente):
        """Gera documento Word personalizado substituindo os placeholders"""
        inicio = time.time()
        doc = Document(self.caminho_template)
        
        # Substitui placeholders PRESERVANDO a formatação original
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                texto_run = run.text
                for key, value in dados_cliente.items():
                    if value:
                        placeholder = f"{{{key}}}"
                        if placeholder in texto_run:
                            run.text = texto_run.replace(placeholder, str(value))
        
        # Também verifica tabelas (se houver)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            texto_run = run.text
                            for key, value in dados_cliente.items():
                                if value:
                                    placeholder = f"{{{key}}}"
                                    if placeholder in texto_run:
                                        run.text = texto_run.replace(placeholder, str(value))
        
        # Cria nome do arquivo: Proposta_001_2025_Maria_Silva.docx
        nome_cliente = str(dados_cliente['NOME_CLIENTE']).replace(' ', '_')
        id_proposta = str(dados_cliente['ID_PROPOSTA']).replace('/', '_')
        nome_arquivo = f"Proposta_{id_proposta}_{nome_cliente}.docx"
        caminho_saida = os.path.join(self.pasta_saida, nome_arquivo)
        
        doc.save(caminho_saida)
        duracao = time.time() - inicio
        self.logger.sucesso(f"Documento gerado para {dados_cliente['NOME_CLIENTE']} em {duracao:.2f}s")
        return caminho_saida
    
    def processar_clientes(self, limite=None):
        """Método principal que coordena todo o processo de geração"""
        inicio_total = time.time()
        self.logger.inicio_processo("Processamento de clientes")
        
        # Lê dados da planilha
        dados_clientes = self.ler_dados_completos()
        
        if limite:
            dados_clientes = dados_clientes[:limite]
            self.logger.info(f"Limitando processamento a {limite} clientes")
        
        # Valida os dados antes de gerar documentos
        self.logger.info("Iniciando validação de dados")
        validador = ValidadorDados()
        print("Validando dados...")
        
        clientes_validos = []
        for i, cliente in enumerate(dados_clientes, 1):
            if validador.validar_cliente(cliente, i + 1):  # +1 por causa do cabeçalho
                clientes_validos.append(cliente)
        
        relatorio_validacao = validador.obter_relatorio_erros()
        print(relatorio_validacao)
        
        invalidos = len(dados_clientes) - len(clientes_validos)
        self.logger.info(f"Validação concluída: {len(clientes_validos)} válidos, {invalidos} inválidos")
        
        if not clientes_validos:
            self.logger.erro("Nenhum cliente válido encontrado")
            print("Nenhum cliente valido encontrado!")
            return []
        
        # Gera documentos para clientes válidos
        self.logger.info(f"Iniciando geração de {len(clientes_validos)} documentos")
        documentos_gerados = []
        erros_geracao = 0
        print(f"Gerando documentos para {len(clientes_validos)} clientes validos...")
        
        for i, cliente in enumerate(clientes_validos, 1):
            try:
                caminho_doc = self.gerar_documento(cliente)
                documentos_gerados.append(caminho_doc)
                print(f"{i:2d}. {cliente['NOME_CLIENTE']} - OK")
            except Exception as e:
                erros_geracao += 1
                self.logger.erro(f"Falha ao gerar documento para {cliente['NOME_CLIENTE']}: {e}")
                print(f"{i:2d}. {cliente['NOME_CLIENTE']} - ERRO: {e}")
        
        # Estatísticas finais
        duracao_total = time.time() - inicio_total
        sucessos = len(documentos_gerados)
        self.logger.estatisticas(len(clientes_validos), sucessos, erros_geracao)
        self.logger.fim_processo("Processamento de clientes", duracao_total)
        
        return documentos_gerados

# Executa apenas quando o arquivo é rodado diretamente
if __name__ == "__main__":
    planilha = "../data/RPONTES.xlsx"
    template = "../templates/modelo_proposta.docx"
    
    gerador = GeradorRPontes(planilha, template)
    
    print("Iniciando geração de documentos...")
    documentos = gerador.processar_clientes()
    print(f"\nConcluído! {len(documentos)} documentos gerados.")